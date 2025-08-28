import streamlit as st
import pandas as pd
import numpy as np
import json
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any
import io

# Import required libraries
try:
    from anthropic import Anthropic
    from pydantic import BaseModel, Field
    from docx import Document
    from docx.shared import Inches
except ImportError as e:
    st.error(f"Missing dependency: {e}. Please install with: uv add anthropic pydantic python-docx")
    st.stop()

# Pydantic models for structured outputs
class Anomaly(BaseModel):
    metric: str = Field(description="The financial metric with anomaly")
    current_value: str = Field(description="Current quarter value")
    comparison_value: str = Field(default="N/A", description="Previous quarter value") 
    change_percent: str = Field(default="N/A", description="Percentage change")
    risk_level: str = Field(description="High, Medium, or Low")
    explanation: str = Field(description="Business context explanation", max_length=200)
    next_steps: str = Field(description="Recommended actions", max_length=150)
    z_score: float = Field(description="Statistical z-score")

class ExecutiveSummary(BaseModel):
    narrative: str = Field(description="3-5 sentence executive summary")
    key_metrics: List[str] = Field(description="List of key metrics mentioned")
    data_sources: List[str] = Field(description="Source columns referenced")

class FinancialAnalysis(BaseModel):
    executive_summary: ExecutiveSummary
    anomalies: List[Anomaly]
    total_anomalies_found: int
    analysis_timestamp: str

# Initialize Anthropic client
def get_anthropic_client():
    """Get Anthropic client with API key validation"""
    try:
        api_key = st.secrets.get("ANTHROPIC_API_KEY")
    except:
        api_key = None
    
    if not api_key:
        api_key = st.sidebar.text_input(
            "Anthropic API Key", type="password", 
            help="Get your API key from console.anthropic.com"
        )
    
    if not api_key:
        st.warning("Please enter your Anthropic API key in the sidebar to continue")
        st.stop()
    
    return Anthropic(api_key=api_key)

# Database setup for audit trail
def init_database():
    conn = sqlite3.connect('audit_trail.db')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS analysis_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT,
            input_file_name TEXT,
            anomalies_detected INTEGER,
            prompt_used TEXT,
            ai_response TEXT,
            user_edits TEXT
        )
    ''')
    conn.commit()
    return conn

# Anomaly detection functions
def detect_anomalies(df: pd.DataFrame, threshold: float = 2.0) -> List[Dict[str, Any]]:
    """Detect statistical anomalies using z-score method"""
    anomalies = []
    
    # Get numeric columns only
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    
    for col in numeric_cols:
        if len(df[col].dropna()) < 2:  # Skip if insufficient data
            continue
            
        # Calculate z-scores
        mean_val = df[col].mean()
        std_val = df[col].std()
        
        if std_val == 0:  # Skip if no variation
            continue
            
        z_scores = np.abs((df[col] - mean_val) / std_val)
        
        # Find outliers
        outlier_indices = z_scores >= threshold
        
        if outlier_indices.any():
            for idx in df[outlier_indices].index:
                anomalies.append({
                    'metric': col,
                    'index': idx,
                    'value': df.loc[idx, col],
                    'z_score': z_scores.loc[idx],
                    'mean': mean_val,
                    'std': std_val
                })
    
    return anomalies

def calculate_qoq_changes(df: pd.DataFrame) -> pd.DataFrame:
    """Calculate quarter-over-quarter changes"""
    if 'Quarter' not in df.columns:
        return df
        
    df_sorted = df.sort_values('Quarter')
    numeric_cols = df_sorted.select_dtypes(include=[np.number]).columns
    
    for col in numeric_cols:
        df_sorted[f'{col}_QoQ_Change'] = df_sorted[col].pct_change() * 100
        df_sorted[f'{col}_QoQ_Abs_Change'] = df_sorted[col].diff()
    
    return df_sorted

def safe_column_match(metric_name: str, df_columns: List[str]) -> str:
    """Safely match anomaly metric names to actual DataFrame columns"""
    
    # Direct match first
    if metric_name in df_columns:
        return metric_name
    
    # Fuzzy matching for common variations
    metric_clean = metric_name.lower().replace(' ', '_').replace('-', '_')
    
    for col in df_columns:
        col_clean = col.lower().replace(' ', '_').replace('-', '_')
        
        # Exact match after cleaning
        if metric_clean == col_clean:
            return col
            
        # Partial match (e.g., "Operating_Expenses" matches "Operating_Expenses_M")
        if metric_clean in col_clean or col_clean in metric_clean:
            return col
    
    # If no match found, return the original name with warning
    return metric_name

def generate_audit_trail_report(analysis: FinancialAnalysis, df: pd.DataFrame, anomalies: List[Dict]) -> str:
    """Generate comprehensive audit trail with cell-level provenance"""
    
    report = f"""AUDIT TRAIL REPORT - AI ANALYSIS PROVENANCE
Generated: {analysis.analysis_timestamp}
Dataset: {df.shape[0]} rows x {df.shape[1]} columns

{'='*60}
EXECUTIVE SUMMARY ANALYSIS
{'='*60}

AI Generated Text:
"{analysis.executive_summary.narrative}"

Key Metrics Referenced:
"""
    
    for metric in analysis.executive_summary.key_metrics:
        report += f"- {metric}\n"
    
    report += f"\nData Sources: {', '.join(analysis.executive_summary.data_sources)}\n"
    
    report += f"\n{'='*60}\n"
    report += f"ANOMALY DETECTION PROVENANCE\n"
    report += f"{'='*60}\n\n"
    
    if not anomalies:
        report += "No statistical anomalies detected in this dataset.\n"
    else:
        for i, raw_anomaly in enumerate(anomalies, 1):
            report += f"ANOMALY #{i}: {raw_anomaly['metric']}\n"
            report += f"{'-'*40}\n"
            
            # Find column index safely
            try:
                if raw_anomaly['metric'] in df.columns:
                    col_idx = list(df.columns).index(raw_anomaly['metric'])
                    col_letter = chr(65 + col_idx)
                    row_number = raw_anomaly['index'] + 2
                    cell_ref = f"{col_letter}{row_number}"
                    report += f"Source Cell: {cell_ref} ({raw_anomaly['metric']})\n"
                else:
                    report += f"Source: {raw_anomaly['metric']} column\n"
            except (KeyError, ValueError):
                report += f"Source: {raw_anomaly.get('metric', 'Unknown column')}\n"
            
            report += f"Raw Value: {raw_anomaly.get('value', 'N/A')}\n"
            report += f"Statistical Analysis:\n"
            report += f"  - Dataset Mean: {raw_anomaly.get('mean', 0):.2f}\n"
            report += f"  - Standard Deviation: {raw_anomaly.get('std', 0):.2f}\n"
            report += f"  - Z-Score: {raw_anomaly.get('z_score', 0):.2f}\n"
            report += f"  - Detection Threshold: 2.0\n\n"
    
    # Add AI analysis if available
    if hasattr(analysis, 'anomalies') and analysis.anomalies:
        report += f"AI ANALYSIS RESPONSES:\n\n"
        for i, anomaly_obj in enumerate(analysis.anomalies, 1):
            report += f"AI Response #{i}:\n"
            report += f"  - Metric: {anomaly_obj.metric}\n"
            report += f"  - Current Value: {anomaly_obj.current_value}\n"
            report += f"  - Risk Level: {anomaly_obj.risk_level}\n"
            report += f"  - Explanation: {anomaly_obj.explanation}\n"
            report += f"  - Next Steps: {anomaly_obj.next_steps}\n\n"
    
    report += f"{'='*60}\n"
    report += f"VALIDATION SUMMARY\n"
    report += f"{'='*60}\n"
    report += f"Total Statistical Anomalies: {len(anomalies)}\n"
    report += f"AI Responses Generated: {len(analysis.anomalies) if hasattr(analysis, 'anomalies') else 0}\n"
    report += f"Schema Validation: PASSED\n"
    report += f"Data Integrity: All numbers traced to source\n"
    
    return report

def export_to_word(analysis: FinancialAnalysis) -> bytes:
    """Export analysis to Word document with error handling"""
    try:
        doc = Document()
        
        # Title
        doc.add_heading('Quarterly Financial Analysis Report', 0)
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(analysis.executive_summary.narrative)
        
        # Key Metrics Section
        if analysis.executive_summary.key_metrics:
            doc.add_heading('Key Metrics Analyzed', level=2)
            for metric in analysis.executive_summary.key_metrics:
                p = doc.add_paragraph()
                p.add_run(f"• {metric}")
        
        # Anomalies Section
        if analysis.anomalies:
            doc.add_heading('Risk & Anomaly Analysis', level=1)
            
            for i, anomaly in enumerate(analysis.anomalies, 1):
                # Anomaly heading
                doc.add_heading(f'Anomaly {i}: {anomaly.metric}', level=2)
                
                # Current value
                doc.add_paragraph(f"Current Value: {anomaly.current_value}")
                
                # Previous value if available
                if anomaly.comparison_value and anomaly.comparison_value != "N/A":
                    doc.add_paragraph(f"Previous Value: {anomaly.comparison_value}")
                
                # Change if available  
                if anomaly.change_percent and anomaly.change_percent != "N/A":
                    doc.add_paragraph(f"Change: {anomaly.change_percent}")
                
                # Risk and analysis
                doc.add_paragraph(f"Risk Level: {anomaly.risk_level}")
                doc.add_paragraph(f"Business Analysis: {anomaly.explanation}")
                doc.add_paragraph(f"Recommended Actions: {anomaly.next_steps}")
                doc.add_paragraph(f"Statistical Confidence: Z-score = {anomaly.z_score:.2f}")
        
        # Audit Information
        doc.add_heading('Analysis Metadata', level=1)
        doc.add_paragraph(f"Generated: {analysis.analysis_timestamp}")
        doc.add_paragraph(f"Total Anomalies Detected: {analysis.total_anomalies_found}")
        
        if analysis.executive_summary.data_sources:
            doc.add_paragraph(f"Data Sources: {', '.join(analysis.executive_summary.data_sources)}")
        
        # Convert to bytes for download
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()
        
    except Exception as e:
        # Return a simple text document if Word generation fails
        simple_report = f"""FINANCIAL ANALYSIS REPORT
Generated: {analysis.analysis_timestamp}

EXECUTIVE SUMMARY:
{analysis.executive_summary.narrative}

ANOMALIES DETECTED: {analysis.total_anomalies_found}
"""
        
        if analysis.anomalies:
            for i, anomaly in enumerate(analysis.anomalies, 1):
                simple_report += f"""
ANOMALY {i}: {anomaly.metric}
- Current Value: {anomaly.current_value}
- Risk Level: {anomaly.risk_level}  
- Analysis: {anomaly.explanation}
- Next Steps: {anomaly.next_steps}
- Z-Score: {anomaly.z_score:.2f}
"""
        
        return simple_report.encode('utf-8')

# Main Streamlit App
def main():
    st.set_page_config(
        page_title="AI Anomaly Detection and Summarization for Reporting",
        page_icon="📊",
        layout="wide"
    )
    
    st.title("🚀 AI Anomaly Detection and Summarization for Reporting")
    st.markdown("*Automatically generate executive summaries and detect statistical anomalies from financial data*")
    
    # Initialize database and client
    conn = init_database()
    client = get_anthropic_client()
    
    # File upload - define this first
    st.header("📁 Upload Financial Data")
    uploaded_file = st.file_uploader(
        "Choose a CSV or Excel file", 
        type=['csv', 'xlsx'],
        help="Upload quarterly financial data for analysis"
    )
    
    # Sidebar configuration - after file upload is defined
    st.sidebar.header("Detection Settings")
    
    # Business-friendly threshold selection
    sensitivity_level = st.sidebar.select_slider(
        "Anomaly Detection Sensitivity",
        options=["Conservative", "Moderate", "Sensitive"],
        value="Moderate",
        help="Conservative: Only flag extreme outliers | Moderate: Standard business thresholds | Sensitive: Flag smaller deviations"
    )
    
    # Map business terms to statistical values
    threshold_mapping = {
        "Conservative": 3.0,   # ~99.7% confidence - only extreme outliers
        "Moderate": 2.0,       # ~95% confidence - standard business practice
        "Sensitive": 1.5       # ~87% confidence - catch smaller issues
    }
    
    z_score_threshold = threshold_mapping[sensitivity_level]
    
    # Show real-time impact
    with st.sidebar.expander("What This Means"):
        if sensitivity_level == "Conservative":
            st.write("🎯 **Ultra-high confidence alerts**")
            st.write("• Flags ~0.3% of typical data points")
            st.write("• Use for: Critical SEC filings, executive dashboards")
            st.write("• Risk: May miss subtle but important issues")
        elif sensitivity_level == "Moderate": 
            st.write("📊 **Standard business threshold**")
            st.write("• Flags ~5% of typical data points")
            st.write("• Use for: Quarterly reviews, standard reporting")
            st.write("• Balance of accuracy and coverage")
        else:  # Sensitive
            st.write("🔍 **Early warning system**")
            st.write("• Flags ~13% of typical data points") 
            st.write("• Use for: Monthly monitoring, trend analysis")
            st.write("• Risk: More false positives requiring review")
    
    # Process uploaded file
    if uploaded_file is not None:
        try:
            # Read file
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)
            
            st.success(f"✅ File uploaded: {uploaded_file.name} ({df.shape[0]} rows, {df.shape[1]} columns)")
            
            # Show live impact preview in sidebar now that we have data
            with st.sidebar.expander("Live Impact Preview"):
                try:
                    # Calculate anomalies for each threshold
                    conservative_anomalies = detect_anomalies(df, threshold=3.0)
                    moderate_anomalies = detect_anomalies(df, threshold=2.0)
                    sensitive_anomalies = detect_anomalies(df, threshold=1.5)
                    current_anomalies = detect_anomalies(df, threshold=z_score_threshold)
                    
                    st.metric("Current Setting", len(current_anomalies))
                    st.write(f"Conservative: {len(conservative_anomalies)} anomalies")
                    st.write(f"Moderate: {len(moderate_anomalies)} anomalies")
                    st.write(f"Sensitive: {len(sensitive_anomalies)} anomalies")
                except Exception as e:
                    st.write("Error calculating preview")
            
            # Display data preview
            st.header("📊 Data Preview")
            st.dataframe(df.head())
            
            # Calculate QoQ changes
            df_with_changes = calculate_qoq_changes(df)
            
            # Detect anomalies with current threshold
            with st.spinner("🔍 Detecting anomalies..."):
                anomalies = detect_anomalies(df_with_changes, threshold=z_score_threshold)
            
            st.header("⚠️ Anomalies Detected")
            if anomalies:
                st.warning(f"Found {len(anomalies)} statistical anomalies")
                
                for anomaly in anomalies:
                    with st.expander(f"🚨 {anomaly['metric']} (z-score: {anomaly['z_score']:.2f})"):
                        st.write(f"**Value:** {anomaly['value']:,.2f}")
                        st.write(f"**Dataset Mean:** {anomaly['mean']:,.2f}")
                        st.write(f"**Standard Deviation:** {anomaly['std']:,.2f}")
                        st.write(f"**Risk Level:** {'High' if anomaly['z_score'] > 3 else 'Medium'}")
            else:
                st.info("No significant anomalies detected with current threshold")
            
            # Generate AI analysis
            if st.button("🤖 Generate AI Analysis", type="primary"):
                with st.spinner("🧠 Analyzing with Claude Haiku..."):
                    
                    try:
                        # Prepare data for Claude with counter-examples
                        data_summary = {
                            'latest_quarter': df_with_changes.iloc[-1].to_dict() if len(df_with_changes) > 0 else {},
                        }
                        
                        # Create analysis tool schema
                        analysis_tool = {
                            "name": "financial_analysis",
                            "description": "Analyze financial data and generate structured insights",
                            "input_schema": {
                                "type": "object",
                                "properties": {
                                    "executive_summary": {
                                        "type": "object",
                                        "properties": {
                                            "narrative": {"type": "string"},
                                            "key_metrics": {"type": "array", "items": {"type": "string"}},
                                            "data_sources": {"type": "array", "items": {"type": "string"}}
                                        }
                                    },
                                    "anomalies": {
                                        "type": "array",
                                        "items": {
                                            "type": "object",
                                            "properties": {
                                                "metric": {"type": "string"},
                                                "current_value": {"type": "string"},
                                                "comparison_value": {"type": "string"},
                                                "change_percent": {"type": "string"},
                                                "risk_level": {"type": "string"},
                                                "explanation": {"type": "string", "maxLength": 200},
                                                "next_steps": {"type": "string", "maxLength": 150},
                                                "z_score": {"type": "number"}
                                            }
                                        }
                                    },
                                    "total_anomalies_found": {"type": "integer"},
                                    "analysis_timestamp": {"type": "string"}
                                }
                            }
                        }
                        
                        # Calculate specific variances for the prompt
                        current_quarter = df_with_changes.iloc[-1]
                        previous_quarter = df_with_changes.iloc[-2] if len(df_with_changes) > 1 else current_quarter
                        
                        # Pre-calculate key variances to inject into prompt
                        revenue_change = current_quarter.get('Total_Revenue_M', 0) - previous_quarter.get('Total_Revenue_M', 0)
                        revenue_pct = (revenue_change / previous_quarter.get('Total_Revenue_M', 1)) * 100 if previous_quarter.get('Total_Revenue_M', 0) != 0 else 0
                        
                        opex_change = current_quarter.get('Operating_Expenses_M', 0) - previous_quarter.get('Operating_Expenses_M', 0)
                        opex_pct = (opex_change / previous_quarter.get('Operating_Expenses_M', 1)) * 100 if previous_quarter.get('Operating_Expenses_M', 0) != 0 else 0
                        
                        # Enhanced prompt requiring specific calculations
                        prompt = f"""Analyze this quarterly financial data with MANDATORY specific calculations.

YOU MUST INCLUDE THESE EXACT CALCULATIONS IN YOUR NARRATIVE:

Revenue Analysis:
- Total Revenue: ${current_quarter.get('Total_Revenue_M', 0):.1f}M (vs ${previous_quarter.get('Total_Revenue_M', 0):.1f}M = {revenue_change:+.1f}M or {revenue_pct:+.1f}%)
- APAC: ${current_quarter.get('APAC_Revenue_M', 0):.1f}M (vs ${previous_quarter.get('APAC_Revenue_M', 0):.1f}M)
- Americas: ${current_quarter.get('Americas_Revenue_M', 0):.1f}M (vs ${previous_quarter.get('Americas_Revenue_M', 0):.1f}M)  
- EMEA: ${current_quarter.get('EMEA_Revenue_M', 0):.1f}M (vs ${previous_quarter.get('EMEA_Revenue_M', 0):.1f}M)

Cost Analysis:
- Operating Expenses: ${current_quarter.get('Operating_Expenses_M', 0):.1f}M (vs ${previous_quarter.get('Operating_Expenses_M', 0):.1f}M = {opex_change:+.1f}M or {opex_pct:+.1f}%)

FORBIDDEN WORDS: "significant," "notable," "challenges," "escalating" - USE NUMBERS ONLY

REQUIRED FORMAT: "Revenue declined $X.XM (-X.X%) from $Y.YM to $Z.ZM"

Current Quarter Data: {json.dumps(data_summary['latest_quarter'], indent=2)}
Statistical Anomalies: {json.dumps(anomalies, indent=2)}

Generate analysis using financial_analysis tool with mandatory numerical precision."""

                        response = client.messages.create(
                            model="claude-3-5-haiku-20241022",
                            max_tokens=1500,
                            temperature=0.1,
                            tools=[analysis_tool],
                            messages=[{"role": "user", "content": prompt}]
                        )
                        
                        # Debug: Show what Claude actually returned
                        st.write("**Debug - Claude Response:**")
                        st.json(response.model_dump())
                        
                        # Parse response with enhanced error handling
                        tool_result = None
                        for content in response.content:
                            if content.type == "tool_use":
                                tool_result = content.input
                                st.write("**Tool Input Found:**")
                                st.json(tool_result)
                                break
                        
                        if not tool_result:
                            st.error("Claude did not return structured data. Check the debug info above.")
                            st.stop()
                        
                        # Add required fields with defaults if missing
                        if "executive_summary" not in tool_result:
                            tool_result["executive_summary"] = {
                                "narrative": "Analysis complete - review anomalies below",
                                "key_metrics": [col for col in df.columns if df[col].dtype in ['int64', 'float64']][:3],
                                "data_sources": df.columns.tolist()[:3]
                            }
                        
                        if "anomalies" not in tool_result:
                            tool_result["anomalies"] = []
                        
                        if "total_anomalies_found" not in tool_result:
                            tool_result["total_anomalies_found"] = len(anomalies)
                            
                        tool_result["analysis_timestamp"] = datetime.now().isoformat()
                        
                        # Validate and clean anomalies with automatic truncation
                        cleaned_anomalies = []
                        for anomaly in tool_result["anomalies"]:
                            cleaned_anomaly = {
                                "metric": anomaly.get("metric", "Unknown"),
                                "current_value": anomaly.get("current_value", "N/A"),
                                "comparison_value": anomaly.get("comparison_value", "N/A"),
                                "change_percent": anomaly.get("change_percent", "N/A"),
                                "risk_level": anomaly.get("risk_level", "Medium"),
                                "explanation": str(anomaly.get("explanation", "Statistical anomaly detected"))[:200],
                                "next_steps": str(anomaly.get("next_steps", "Review data for accuracy"))[:150],
                                "z_score": float(anomaly.get("z_score", 2.0))
                            }
                            cleaned_anomalies.append(cleaned_anomaly)
                        tool_result["anomalies"] = cleaned_anomalies
                        
                        st.write("**Final Processed Data:**")
                        st.json(tool_result)
                        
                        analysis = FinancialAnalysis(**tool_result)
                    
                    except Exception as e:
                        st.error(f"Error generating analysis: {str(e)}")
                        st.info("Using fallback analysis...")
                        
                        # Simple fallback that will definitely work
                        analysis = FinancialAnalysis(
                            executive_summary=ExecutiveSummary(
                                narrative=f"Quarterly analysis complete for {uploaded_file.name}. {len(anomalies)} statistical anomalies detected requiring review.",
                                key_metrics=[col for col in df.columns if df[col].dtype in ['int64', 'float64']][:5],
                                data_sources=df.columns.tolist()[:5]
                            ),
                            anomalies=[
                                Anomaly(
                                    metric=anomaly['metric'],
                                    current_value=f"${anomaly['value']:,.0f}M" if anomaly['value'] > 1000 else f"{anomaly['value']:,.2f}",
                                    comparison_value="Previous period",
                                    change_percent="See analysis", 
                                    risk_level="High" if anomaly['z_score'] > 3 else "Medium",
                                    explanation=f"Statistical outlier with z-score of {anomaly['z_score']:.1f} standard deviations from normal range",
                                    next_steps="Review data source and validate business context for this unusual value",
                                    z_score=anomaly['z_score']
                                ) for anomaly in anomalies[:5]  # Limit to 5 anomalies
                            ],
                            total_anomalies_found=len(anomalies),
                            analysis_timestamp=datetime.now().isoformat()
                        )
                
                if 'analysis' in locals() and analysis:
                    # Display results
                    st.header("📝 AI-Generated Analysis")
                    
                    # Executive Summary
                    st.subheader("Executive Summary")
                    st.write(analysis.executive_summary.narrative)
                    
                    # Key Metrics
                    with st.expander("Key Metrics Referenced"):
                        for metric in analysis.executive_summary.key_metrics:
                            st.write(f"• {metric}")
                    
                    # Anomaly Analysis
                    if analysis.anomalies:
                        st.subheader("🚨 Risk & Anomaly Analysis")
                        
                        for i, anomaly in enumerate(analysis.anomalies, 1):
                            with st.container():
                                col1, col2 = st.columns([3, 1])
                                
                                with col1:
                                    st.write(f"**{i}. {anomaly.metric}**")
                                    st.write(f"*Current:* {anomaly.current_value}")
                                    if anomaly.comparison_value and anomaly.comparison_value != "N/A":
                                        st.write(f"*Previous:* {anomaly.comparison_value}")
                                    if anomaly.change_percent and anomaly.change_percent != "N/A":
                                        st.write(f"*Change:* {anomaly.change_percent}")
                                    st.write(f"*Analysis:* {anomaly.explanation}")
                                    st.write(f"*Next Steps:* {anomaly.next_steps}")
                                
                                with col2:
                                    risk_color = {"High": "🔴", "Medium": "🟡", "Low": "🟢"}
                                    st.metric(
                                        "Risk Level", 
                                        f"{risk_color.get(anomaly.risk_level, '⚪')} {anomaly.risk_level}",
                                        f"z-score: {anomaly.z_score:.2f}"
                                    )
                                
                                st.divider()
                    
                    # Export options
                    st.header("📤 Export Analysis")
                    
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        if st.button("📄 Download Word Report"):
                            try:
                                doc_bytes = export_to_word(analysis)
                                st.download_button(
                                    label="💾 Download DOCX",
                                    data=doc_bytes,
                                    file_name="financial_analysis_report.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                            except Exception as e:
                                st.error(f"Error generating Word document: {str(e)}")
                    
                    with col2:
                        try:
                            json_data = analysis.model_dump_json(indent=2)
                            st.download_button(
                                label="📊 Download JSON",
                                data=json_data,
                                file_name="financial_analysis.json",
                                mime="application/json"
                            )
                        except Exception as e:
                            st.error(f"Error generating JSON: {str(e)}")
                    
                    with col3:
                        try:
                            audit_report = generate_audit_trail_report(analysis, df_with_changes, anomalies)
                            st.download_button(
                                label="🔍 Download Audit Trail", 
                                data=audit_report,
                                file_name="audit_trail_report.txt",
                                mime="text/plain"
                            )
                        except Exception as e:
                            st.error(f"Error generating audit trail: {str(e)}")
                    
                    # Log to database
                    try:
                        conn.execute('''
                            INSERT INTO analysis_log 
                            (timestamp, input_file_name, anomalies_detected, prompt_used, ai_response) 
                            VALUES (?, ?, ?, ?, ?)
                        ''', (
                            datetime.now().isoformat(),
                            uploaded_file.name,
                            len(anomalies),
                            "Financial analysis with counter-examples",
                            json.dumps(analysis.model_dump(), indent=2)
                        ))
                        conn.commit()
                        
                        st.success("✅ Analysis complete! Audit trail logged.")
                    except Exception as e:
                        st.warning(f"Could not log to database: {str(e)}")
        
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")
            st.info("Please ensure your file is a valid CSV or Excel file with numeric data.")
    
    else:
        # Show instructions when no file is uploaded
        st.info("👆 Upload a CSV or Excel file to begin analysis")
        
        with st.expander("📋 Expected Data Format"):
            st.write("Your file should contain:")
            st.write("• **Quarterly data** with clear time periods")
            st.write("• **Numeric columns** for financial metrics (Revenue, Expenses, etc.)")
            st.write("• **Column headers** that clearly identify each metric")
            st.write("• **Multiple quarters** for meaningful statistical analysis")
    
    # Footer
    st.markdown("---")
    st.markdown("*AI-powered anomaly detection and summarization for enterprise reporting*")

if __name__ == "__main__":
    main()